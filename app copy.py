from flask import Flask, render_template, request, jsonify, send_file, Response
from flask_cors import CORS
from openai import OpenAI
import requests as http_requests
import os
import io
import base64
import json
import PyPDF2
import docx
import fitz  # PyMuPDF
from fpdf import FPDF

app = Flask(__name__)
CORS(app)

# ── Model 1: Text Summarization (dracarys via OpenAI SDK) ──
text_client = OpenAI(
    base_url="https://integrate.api.nvidia.com/v1",
    api_key="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"
)
TEXT_MODEL = "abacusai/dracarys-llama-3.1-70b-instruct"

# ── Model 2: Image/Video Analysis (Kimi K2.5 via raw requests) ──
KIMI_API_URL = "https://integrate.api.nvidia.com/v1/chat/completions"
KIMI_API_KEY = "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"
VISION_MODEL = "moonshotai/kimi-k2.5"

# ── Model 3: Structured Output - Notes/Flowchart/Exam (Nemotron via OpenAI SDK) ──
nemotron_client = OpenAI(
    base_url="https://integrate.api.nvidia.com/v1",
    api_key="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"
)
NEMOTRON_MODEL = "nvidia/llama-3.3-nemotron-super-49b-v1.5"

# ── Mode-specific system prompts ──
MODE_PROMPTS = {
    "notes": """You are an expert study-notes generator. Given the following content, create comprehensive, well-structured study notes. Follow this format:
- Use clear hierarchical headers (# Topic, ## Subtopic, ### Key Point)
- Use bullet points for key facts and definitions
- **Bold** all important terms and definitions
- Add a "Key Takeaways" section at the end
- Keep language concise but thorough
- Use numbered lists for sequential processes
Do NOT add any preamble. Start directly with the notes.""",

    "flowchart": """You are a Mermaid.js diagram expert. Given the following content, create a clear and well-structured Mermaid flowchart that visualizes the key concepts, processes, or relationships.

Rules:
- Output ONLY the raw Mermaid code inside a ```mermaid code block
- Use `graph TD` (top-down) orientation
- Use descriptive labels in brackets `[Label]` for process nodes
- Use curly braces `{Decision?}` for decision/diamond nodes
- Use `([Label])` for rounded nodes (start/end)
- Connect nodes with labeled arrows where helpful: `A -->|yes| B`
- Use subgraphs to group related concepts if the content has multiple sections
- Keep node labels short but meaningful (max 6 words)
- Ensure the diagram is syntactically valid Mermaid
Do NOT add any explanation before or after the mermaid block.""",

    "exam": """You are an expert exam preparation assistant. Given the following content, create comprehensive exam-ready study material. Structure it as follows:

## 📝 Key Concepts
List the most important concepts with brief explanations.

## ❓ Short Answer Questions
Generate 5-8 short answer questions with model answers.

## 🔘 Multiple Choice Questions
Generate 5-8 MCQs with 4 options each. Mark the correct answer with ✅.

## 📋 Fill in the Blanks
Generate 5 fill-in-the-blank questions with answers.

## 🧠 Mnemonics & Memory Aids
Create helpful mnemonics or memory tricks for the key concepts.

## ⚡ Quick Revision Points
Bullet-point summary of the most exam-critical facts.

Make the content thorough, accurate, and exam-focused. Do NOT add any preamble."""
}

# ── Difficulty level prefixes ──
DIFFICULTY_PREFIX = {
    "simple": "Use simple, easy-to-understand language suitable for beginners. Avoid jargon. Explain like teaching a 10th grader.",
    "intermediate": "Use clear, standard academic language. Balance detail with readability.",
    "advanced": "Use precise, technical language. Include in-depth analysis, edge cases, and expert-level detail."
}

# Supported file extensions
TEXT_EXTS = {'txt'}
DOC_EXTS = {'pdf', 'docx'}
IMAGE_EXTS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
VIDEO_EXTS = {'mp4', 'webm'}
ALL_MEDIA_EXTS = IMAGE_EXTS | VIDEO_EXTS

MIME_MAP = {
    'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
    'gif': 'image/gif', 'webp': 'image/webp',
    'mp4': 'video/mp4', 'webm': 'video/webm',
}


@app.route('/')
def home():
    return render_template('index.html')


@app.route('/summarize', methods=['POST'])
def summarize():
    text_content = ''
    media_parts = []

    # ── Process uploaded files ──
    uploaded_files = request.files.getlist('files') + request.files.getlist('file')
    for file in uploaded_files:
        if file.filename == '':
            continue
        ext = file.filename.rsplit('.', 1)[-1].lower()
        try:
            if ext in TEXT_EXTS:
                text_content += file.read().decode('utf-8') + '\n'
            elif ext == 'pdf':
                pdf_reader = PyPDF2.PdfReader(file)
                extracted_pdf_text = ''.join(page.extract_text() or '' for page in pdf_reader.pages)
                if not extracted_pdf_text.strip():
                    # Fallback for scanned/image-based PDFs: convert pages to images
                    file.seek(0)
                    pdf_bytes = file.read()
                    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    # Limit to first 20 pages to prevent huge payloads, if necessary
                    for i in range(min(len(doc), 20)):
                        pix = doc[i].get_pixmap(dpi=150)
                        b64 = base64.b64encode(pix.tobytes("png")).decode('utf-8')
                        media_parts.append({
                            "type": "image_url",
                            "image_url": {"url": f"data:image/png;base64,{b64}"}
                        })
                else:
                    text_content += extracted_pdf_text + '\n'
            elif ext == 'docx':
                doc = docx.Document(file)
                text_content += '\n'.join(para.text for para in doc.paragraphs) + '\n'
            elif ext in ALL_MEDIA_EXTS:
                raw = file.read()
                b64 = base64.b64encode(raw).decode('utf-8')
                mime = MIME_MAP.get(ext, 'application/octet-stream')
                media_parts.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}"}
                })
            else:
                return jsonify({'error': f'Unsupported file type: .{ext}'}), 400
        except Exception as e:
            return jsonify({'error': f'Failed to process {file.filename}: {str(e)}'}), 500

    # Fallback: form text or JSON text
    if not text_content and 'text' in request.form:
        text_content = request.form['text']
    elif not text_content and request.is_json:
        data = request.json
        text_content = data.get('text', '')

    if not text_content.strip() and not media_parts:
        return jsonify({'error': 'Please provide text, a document, or media to analyze.'}), 400

    # ── Get the output mode and difficulty ──
    mode = request.form.get('mode', 'summarize').strip().lower()
    difficulty = request.form.get('difficulty', 'intermediate').strip().lower()
    has_media = len(media_parts) > 0

    # ── Route to the correct model ──
    if has_media:
        return Response(stream_kimi(text_content.strip(), media_parts), mimetype='text/plain')
    elif mode in ('notes', 'flowchart', 'exam'):
        return Response(stream_nemotron(text_content.strip(), mode, difficulty), mimetype='text/plain')
    else:
        return Response(stream_dracarys(text_content.strip(), difficulty), mimetype='text/plain')


def stream_dracarys(text, difficulty='intermediate'):
    """Stream text summarization via dracarys (OpenAI SDK)."""
    diff_instruction = DIFFICULTY_PREFIX.get(difficulty, DIFFICULTY_PREFIX['intermediate'])
    prompt = f"{diff_instruction}\n\nPlease provide a concise text summary of the following content:\n\n{text}"
    try:
        completion = text_client.chat.completions.create(
            model=TEXT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            top_p=1,
            max_tokens=1024,
            stream=True
        )
        for chunk in completion:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content
    except Exception as e:
        yield f"\n\n[Error: {str(e)}]"


def stream_nemotron(text, mode, difficulty='intermediate'):
    """Stream structured output (notes/flowchart/exam) via Nemotron Super 49B."""
    diff_instruction = DIFFICULTY_PREFIX.get(difficulty, DIFFICULTY_PREFIX['intermediate'])
    system_prompt = f"{diff_instruction}\n\n{MODE_PROMPTS.get(mode, '')}"
    try:
        completion = nemotron_client.chat.completions.create(
            model=NEMOTRON_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            temperature=0.6,
            top_p=0.95,
            max_tokens=16384,
            frequency_penalty=0,
            presence_penalty=0,
            stream=True
        )
        for chunk in completion:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content
    except Exception as e:
        yield f"\n\n[Error: {str(e)}]"


def stream_kimi(text, media_parts):
    """Stream image/video analysis via Kimi K2.5 (raw requests + SSE)."""
    prompt_text = text if text else "Please analyze and describe the attached media in detail."

    content_parts = [{"type": "text", "text": prompt_text}]
    content_parts.extend(media_parts)

    try:
        headers = {
            "Authorization": f"Bearer {KIMI_API_KEY}",
            "Accept": "text/event-stream",
            "Content-Type": "application/json",
        }
        payload = {
            "model": VISION_MODEL,
            "messages": [{"role": "user", "content": content_parts}],
            "max_tokens": 16384,
            "temperature": 1.00,
            "top_p": 1.00,
            "stream": True,
            "chat_template_kwargs": {"thinking": True},
        }

        response = http_requests.post(
            KIMI_API_URL, headers=headers, json=payload,
            stream=True, timeout=120
        )

        if response.status_code != 200:
            yield f"\n\n[API Error {response.status_code}: {response.text[:500]}]"
            return

        in_reasoning = False
        reasoning_done = False

        for line in response.iter_lines():
            if not line:
                continue
            decoded = line.decode('utf-8')
            if decoded.startswith('data: '):
                data_str = decoded[6:]
                if data_str.strip() == '[DONE]':
                    break
                try:
                    chunk = json.loads(data_str)
                    choices = chunk.get('choices', [])
                    if choices:
                        delta = choices[0].get('delta', {})
                        content = delta.get('content')
                        reasoning = delta.get('reasoning_content')
                        if reasoning:
                            if not in_reasoning:
                                yield "> **Thinking Process...**\n> "
                                in_reasoning = True
                            yield reasoning.replace('\n', '\n> ')
                        
                        if content is not None:
                            if in_reasoning and not reasoning_done:
                                yield "\n\n"
                                reasoning_done = True
                            yield content
                except json.JSONDecodeError:
                    continue

    except http_requests.exceptions.Timeout:
        yield "\n\n[Error: Request timed out.]"
    except Exception as e:
        yield f"\n\n[Error: {str(e)}]"


@app.route('/export_docx', methods=['POST'])
def export_docx():
    data = request.json
    text = data.get('text', '')
    doc = docx.Document()
    doc.add_heading('AI Text Summary', 0)
    doc.add_paragraph(text)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(
        file_stream, as_attachment=True, download_name='summary.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/export_pdf', methods=['POST'])
def export_pdf():
    data = request.json
    text = data.get('text', '')

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    
    safe_text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(w=0, h=6, text=safe_text)
    
    pdf_bytes = pdf.output()
    file_stream = io.BytesIO(pdf_bytes)
    file_stream.seek(0)
    return send_file(
        file_stream, as_attachment=True, download_name='summary.pdf',
        mimetype='application/pdf'
    )


@app.route('/export_image_pdf', methods=['POST'])
def export_image_pdf():
    """Convert an uploaded image (flowchart PNG) to a PDF."""
    file = request.files.get('image')
    if not file:
        return jsonify({'error': 'No image provided'}), 400

    img_bytes = file.read()

    pdf = FPDF(orientation='L')  # Landscape for flowcharts
    pdf.add_page()

    # Save temp image to embed in PDF
    import tempfile, os
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    tmp.write(img_bytes)
    tmp.close()

    try:
        # Fit image to page with margins
        pdf.image(tmp.name, x=10, y=10, w=pdf.w - 20)
    except Exception:
        pdf.set_font("Helvetica", size=14)
        pdf.cell(w=0, h=10, text="Flowchart image could not be embedded.")
    finally:
        os.unlink(tmp.name)

    pdf_bytes = pdf.output()
    file_stream = io.BytesIO(pdf_bytes)
    file_stream.seek(0)
    return send_file(
        file_stream, as_attachment=True, download_name='flowchart.pdf',
        mimetype='application/pdf'
    )


if __name__ == '__main__':
    app.run(debug=True, port=5000)
