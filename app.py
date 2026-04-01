from flask import Flask, render_template, request, jsonify, send_file, Response
from flask_cors import CORS
from openai import OpenAI
import requests as http_requests
import os
import io
import base64
import json
import PyPDF2
import docx

app = Flask(__name__)
CORS(app)

# ── Model 1: Text Summarization (dracarys via OpenAI SDK) ──
text_client = OpenAI(
    base_url="https://integrate.api.nvidia.com/v1",
    api_key="Enter your API here"
)
TEXT_MODEL = "abacusai/dracarys-llama-3.1-70b-instruct"

# ── Model 2: Image/Video Analysis (Kimi K2.5 via raw requests) ──
KIMI_API_URL = "https://integrate.api.nvidia.com/v1/chat/completions"
KIMI_API_KEY = "Enter you API here"
VISION_MODEL = "moonshotai/kimi-k2.5"

# Supported file extensions
TEXT_EXTS = {'txt'}
DOC_EXTS = {'pdf', 'docx'}
IMAGE_EXTS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
VIDEO_EXTS = {'mp4', 'webm'}
ALL_MEDIA_EXTS = IMAGE_EXTS | VIDEO_EXTS

MIME_MAP = {
    'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
    'gif': 'image/gif', 'webp': 'image/webp',
    'mp4': 'video/mp4', 'webm': 'video/webm',
}


@app.route('/')
def home():
    return render_template('index.html')


@app.route('/summarize', methods=['POST'])
def summarize():
    text_content = ''
    media_parts = []

    # ── Process uploaded files ──
    uploaded_files = request.files.getlist('files')
    for file in uploaded_files:
        if file.filename == '':
            continue
        ext = file.filename.rsplit('.', 1)[-1].lower()
        try:
            if ext in TEXT_EXTS:
                text_content += file.read().decode('utf-8') + '\n'
            elif ext == 'pdf':
                pdf_reader = PyPDF2.PdfReader(file)
                text_content += ''.join(page.extract_text() or '' for page in pdf_reader.pages) + '\n'
            elif ext == 'docx':
                doc = docx.Document(file)
                text_content += '\n'.join(para.text for para in doc.paragraphs) + '\n'
            elif ext in ALL_MEDIA_EXTS:
                raw = file.read()
                b64 = base64.b64encode(raw).decode('utf-8')
                mime = MIME_MAP.get(ext, 'application/octet-stream')
                media_parts.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}"}
                })
            else:
                return jsonify({'error': f'Unsupported file type: .{ext}'}), 400
        except Exception as e:
            return jsonify({'error': f'Failed to process {file.filename}: {str(e)}'}), 500

    # Fallback: form text or JSON text
    if not text_content and 'text' in request.form:
        text_content = request.form['text']
    elif not text_content and request.is_json:
        data = request.json
        text_content = data.get('text', '')

    if not text_content.strip() and not media_parts:
        return jsonify({'error': 'Please provide text, a document, or media to analyze.'}), 400

    has_media = len(media_parts) > 0

    # ── Route to the correct model ──
    if has_media:
        return Response(stream_kimi(text_content.strip(), media_parts), mimetype='text/plain')
    else:
        return Response(stream_dracarys(text_content.strip()), mimetype='text/plain')


def stream_dracarys(text):
    """Stream text summarization via dracarys (OpenAI SDK)."""
    prompt = f"Please provide a concise text summary of the following content:\n\n{text}"
    try:
        completion = text_client.chat.completions.create(
            model=TEXT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            top_p=1,
            max_tokens=1024,
            stream=True
        )
        for chunk in completion:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content
    except Exception as e:
        yield f"\n\n[Error: {str(e)}]"


def stream_kimi(text, media_parts):
    """Stream image/video analysis via Kimi K2.5 (raw requests + SSE)."""
    prompt_text = text if text else "Please analyze and describe the attached media in detail."

    content_parts = [{"type": "text", "text": prompt_text}]
    content_parts.extend(media_parts)

    try:
        headers = {
            "Authorization": f"Bearer {KIMI_API_KEY}",
            "Accept": "text/event-stream",
            "Content-Type": "application/json",
        }
        payload = {
            "model": VISION_MODEL,
            "messages": [{"role": "user", "content": content_parts}],
            "max_tokens": 16384,
            "temperature": 1.00,
            "top_p": 1.00,
            "stream": True,
            "chat_template_kwargs": {"thinking": True},
        }

        response = http_requests.post(
            KIMI_API_URL, headers=headers, json=payload,
            stream=True, timeout=120
        )

        if response.status_code != 200:
            yield f"\n\n[API Error {response.status_code}: {response.text[:500]}]"
            return

        for line in response.iter_lines():
            if not line:
                continue
            decoded = line.decode('utf-8')
            if decoded.startswith('data: '):
                data_str = decoded[6:]
                if data_str.strip() == '[DONE]':
                    break
                try:
                    chunk = json.loads(data_str)
                    choices = chunk.get('choices', [])
                    if choices:
                        content = choices[0].get('delta', {}).get('content')
                        if content:
                            yield content
                except json.JSONDecodeError:
                    continue

    except http_requests.exceptions.Timeout:
        yield "\n\n[Error: Request timed out.]"
    except Exception as e:
        yield f"\n\n[Error: {str(e)}]"


@app.route('/export_docx', methods=['POST'])
def export_docx():
    data = request.json
    text = data.get('text', '')
    doc = docx.Document()
    doc.add_heading('AI Text Summary', 0)
    doc.add_paragraph(text)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(
        file_stream, as_attachment=True, download_name='summary.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    app.run(debug=True, port=5000)
